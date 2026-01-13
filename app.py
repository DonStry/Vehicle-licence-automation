import io
import json
from datetime import date, datetime
import streamlit as st
from pypdf import PdfReader, PdfWriter
from pypdf.generic import NameObject, BooleanObject, TextStringObject
import base64
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from reportlab.lib.utils import simpleSplit
import re
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from pypdf import PdfReader, PdfWriter



st.set_page_config(page_title="Vehicle Form Assistant", layout="centered")


# =========================
# Session state init
# =========================
if "step" not in st.session_state:
    st.session_state.step = 1
if "user_type" not in st.session_state:
    st.session_state.user_type = "Personal"
if "data" not in st.session_state:
    st.session_state.data = {}
if "generated_pdf_bytes" not in st.session_state:
    st.session_state.generated_pdf_bytes = None
if "generated_pdf_name" not in st.session_state:
    st.session_state.generated_pdf_name = None


# =========================
# Helpers
# =========================

def make_permission_letter_pdf(user_type: str, data: dict) -> bytes:
    """
    Creates a 1-page permission letter PDF (personal or company wording)
    and returns PDF bytes.
    """
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    def draw_text(x, y, text, size=11):
        c.setFont("Helvetica", size)
        c.drawString(x, y, text)

    def draw_wrapped(x, y, text, max_width, line_height=14, size=11):
        c.setFont("Helvetica", size)
        lines = simpleSplit(text, "Helvetica", size, max_width)
        for line in lines:
            c.drawString(x, y, line)
            y -= line_height
        return y

    # Bearer (who will do the transaction)
    bearer_name = data.get("bearer_name", "")
    bearer_surname = data.get("bearer_surname", "")
    bearer_id = data.get("bearer_id", "") or data.get("bearer_id_number", "")

    # Vehicle details (for bottom table)
    lic = data.get("license_number", "")
    reg = data.get("register_number", "")
    vin = data.get("vin", "")
    make = data.get("make", "")
    series = data.get("series_name", "")

    tx = data.get("transaction_type", "")

    # Applicant details + wording
    if user_type == "Personal":
        owner_name = data.get("first_name", "")
        owner_surname = data.get("surname", "")
        owner_id = data.get("id_number", "")

        body = (
            f"I, {owner_name} {owner_surname} with ID number {owner_id}, give permission to "
            f"{bearer_name} {bearer_surname} with ID number {bearer_id} to register/licence "
            f"(as specified under service required) my vehicle/s on my behalf as per vehicle information below. "
            f"I also give permission for the bearer to sign on my behalf if and where needed as well as request "
            f"any information in my name in the event of the transaction not being successful."
        )
    else:
        company_name = data.get("company_name", "")
        ck = data.get("company_ck", "")
        rep_name = data.get("rep_name", "")
        rep_surname = data.get("rep_surname", "")
        rep_id = data.get("rep_id_number", "")

        body = (
            f"I, {rep_name} {rep_surname}, Representative of {company_name} with CK number: {ck}, my ID number {rep_id}, "
            f"give permission to {bearer_name} {bearer_surname} with ID number {bearer_id} to register/licence "
            f"(as specified under service required) my vehicle/s on my behalf as per vehicle information below. "
            f"I also give permission to sign on my behalf if and where needed as well as require any information "
            f"on my name in the event of the transaction not being successful."
        )

    # Page layout
    margin_x = 20 * mm
    y = height - 25 * mm

    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin_x, y, "PERMISSION LETTER")
    y -= 18

    c.setFont("Helvetica", 11)
    y = draw_wrapped(margin_x, y, body, max_width=width - 2 * margin_x, line_height=14, size=11)

    y -= 12
    draw_text(margin_x, y, f"Service required: {tx}", size=11)
    y -= 18

    # Vehicle information table
    c.setFont("Helvetica-Bold", 11)
    c.drawString(margin_x, y, "Vehicle information")
    y -= 12

    table_x = margin_x
    table_w = width - 2 * margin_x
    row_h = 10 * mm

    cols = [
        ("Licence number", lic),
        ("Register number", reg),
        ("VIN", vin),
        ("Make", make),
        ("Series", series),
    ]

    c.setFont("Helvetica", 10)
    for label, value in cols:
        c.rect(table_x, y - row_h, table_w, row_h, stroke=1, fill=0)
        c.drawString(table_x + 4, y - row_h + 7, f"{label}: {value}")
        y -= row_h

    min_y_for_big_signing = 120 * mm
    if y < min_y_for_big_signing:
        c.showPage()
        y = height - 25 * mm

    # BIG GAP after the table before "Signed at"
    y -= 35 * mm

    c.setFont("Helvetica", 11)
    c.drawString(margin_x + 90 * mm, y, "Date: _______________________________")

    # Give a large writing area for signatures
    y -= 25 * mm

    c.drawString(margin_x, y, "Signature (Owner / Representative):")
    c.drawString(margin_x + 100 * mm, y, "Signature (Bearer):")

    # Move the signature LINES down so there is space ABOVE them to sign
    y -= 18 * mm

    c.line(margin_x, y, margin_x + 80 * mm, y)
    c.line(margin_x + 100 * mm, y, margin_x + 180 * mm, y)

    # Extra blank space below signatures
    y -= 20 * mm

    c.showPage()
    c.save()

    buffer.seek(0)
    return buffer.getvalue()


def merge_pdfs(pdf_a: bytes, pdf_b: bytes) -> bytes:
    """
    Returns a single PDF: pdf_a pages first, then pdf_b pages.
    """
    writer = PdfWriter()

    r1 = PdfReader(io.BytesIO(pdf_a))
    r2 = PdfReader(io.BytesIO(pdf_b))

    for p in r1.pages:
        writer.add_page(p)
    for p in r2.pages:
        writer.add_page(p)

    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def pdf_preview(pdf_bytes: bytes, height: int = 900):
    """Render a PDF preview inside Streamlit."""
    if not pdf_bytes:
        return
    b64 = base64.b64encode(pdf_bytes).decode("utf-8")
    html = f"""
        <iframe
            src="data:application/pdf;base64,{b64}"
            width="100%"
            height="{height}"
            style="border:1px solid rgba(255,255,255,0.08); border-radius:10px;"
        ></iframe>
    """
    st.markdown(html, unsafe_allow_html=True)

def make_initials(full_name: str) -> str:
    parts = [p for p in (full_name or "").strip().split() if p]
    return "".join([p[0].upper() for p in parts[:3]])


def make_json_safe(obj):
    if isinstance(obj, dict):
        return {k: make_json_safe(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [make_json_safe(v) for v in obj]
    if isinstance(obj, (date, datetime)):
        return obj.isoformat()
    return obj


def list_pdf_fields(template_path: str) -> list[str]:
    reader = PdfReader(template_path)
    fields = reader.get_fields() or {}
    return sorted(fields.keys())

def generate_permission_letter(data: dict, user_type: str) -> bytes:
    doc = Document()

    doc.add_heading("Permission Letter", level=1)
    doc.add_paragraph("")

    service = data.get("transaction_type", "")

    # Bearer
    bearer_name = f"{data.get('bearer_name', '')} {data.get('bearer_surname', '')}".strip()
    bearer_id = data.get("bearer_id", "")

    if user_type == "Personal":
        client_name = f"{data.get('first_name', '')} {data.get('surname', '')}".strip()
        client_id = data.get("id_number", "")

        paragraph = (
            f"I {client_name} with ID number {client_id} give permission to "
            f"{bearer_name} with ID number {bearer_id} to register/licence "
            f"({service}) my vehicle/s on my behalf as per vehicle information below.\n\n"
            "I also give her permission to sign on my behalf if and where needed, "
            "as well as to request any information in my name in the event of the "
            "transaction not being successful."
        )

    else:
        rep_name = f"{data.get('rep_name', '')} {data.get('rep_surname', '')}".strip()
        rep_id = data.get("rep_id_number", "")
        company_name = data.get("company_name", "")
        company_ck = data.get("company_ck", "")

        paragraph = (
            f"I {rep_name}, representative of {company_name} with CK number {company_ck} "
            f"and my ID number {rep_id}, give permission to {bearer_name} with ID number "
            f"{bearer_id} to register/licence ({service}) my vehicle/s on my behalf as per "
            "vehicle information below.\n\n"
            "I also give her permission to sign on my behalf if and where needed, "
            "as well as to request any information in my name in the event of the "
            "transaction not being successful."
        )

    doc.add_paragraph(paragraph)
    doc.add_paragraph("")

    # Vehicle table
    table = doc.add_table(rows=2, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Make"
    hdr_cells[1].text = "Model"
    hdr_cells[2].text = "Registration"
    hdr_cells[3].text = "Licence"
    hdr_cells[4].text = "VIN"
    hdr_cells[5].text = "Service"

    row = table.rows[1].cells
    row[0].text = data.get("make", "")
    row[1].text = data.get("series_name", "")
    row[2].text = data.get("register_number", "")
    row[3].text = data.get("license_number", "")
    row[4].text = data.get("vin", "")
    row[5].text = service

    doc.add_paragraph("\n\n____________________")

    # Save to bytes
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# =========================
# PDF comb autosize (ALL fields)
# =========================
def autosize_all_combed_fields(writer: PdfWriter, padding: float = 0.62, font_name: str = "Courier"):
    """
    For every widget field that has /MaxLen (comb count),
    compute a font size that fits each character into its box,
    then force a monospaced font + that size.
    """
    for page in writer.pages:
        annots = page.get("/Annots", [])
        if not annots:
            continue

        for annot_ref in annots:
            annot = annot_ref.get_object()
            if annot.get("/Subtype") != "/Widget":
                continue

            # Only fields that have MaxLen can be reliably comb-sized
            maxlen = annot.get("/MaxLen")
            if not maxlen:
                continue

            try:
                maxlen = int(maxlen)
            except Exception:
                continue

            if maxlen <= 0:
                continue

            rect = annot.get("/Rect")
            if not rect or len(rect) < 4:
                continue

            x0, y0, x1, y1 = [float(v) for v in rect]
            width = abs(x1 - x0)
            height = abs(y1 - y0)

            box_w = width / maxlen
            size = int(min(box_w * padding, height * padding))

            if size < 6:
                size = 6

            annot.update({
                NameObject("/DA"): TextStringObject(f"/{font_name} {size} Tf 0 g")
            })

            # remove appearance so the viewer redraws using the DA we set
            if "/AP" in annot:
                del annot["/AP"]

def force_redraw_all_fields(writer: PdfWriter):
    """
    Fixes: fields appear empty until clicked.
    Forces viewers to regenerate appearances by:
    - setting /NeedAppearances = True
    - removing /AP from all widget annotations
    """
    if "/AcroForm" in writer._root_object:
        acroform = writer._root_object["/AcroForm"]
        acroform.update({NameObject("/NeedAppearances"): BooleanObject(True)})

    for page in writer.pages:
        annots = page.get("/Annots", [])
        if not annots:
            continue

        for annot_ref in annots:
            annot = annot_ref.get_object()
            if annot.get("/Subtype") != "/Widget":
                continue

            if "/AP" in annot:
                del annot["/AP"]

def flatten_filled_fields(writer: PdfWriter, field_values: dict, *, default_font="Courier"):
    """
    Draws filled field values onto the PDF pages (flattening appearances),
    so Acrobat/Wondershare show data immediately without clicking fields.

    Works best when fields are combed and have /MaxLen.
    """
    for page in writer.pages:
        annots = page.get("/Annots", [])
        if not annots:
            continue

        # Page size
        mb = page.mediabox
        page_w = float(mb.width)
        page_h = float(mb.height)

        packet = io.BytesIO()
        c = canvas.Canvas(packet, pagesize=(page_w, page_h))

        for annot_ref in annots:
            annot = annot_ref.get_object()
            if annot.get("/Subtype") != "/Widget":
                continue

            t = annot.get("/T")
            if not t:
                continue

            fname = str(t)
            if fname not in field_values:
                continue

            val = str(field_values.get(fname, "") or "")
            if val.strip() == "":
                continue

            rect = annot.get("/Rect")
            if not rect or len(rect) < 4:
                continue

            x0, y0, x1, y1 = [float(v) for v in rect]
            width = abs(x1 - x0)
            height = abs(y1 - y0)

            # Try read font size from /DA if present
            font_size = None
            da = annot.get("/DA")
            if da:
                m = re.search(r"(\d+(?:\.\d+)?)\s+Tf", str(da))
                if m:
                    try:
                        font_size = float(m.group(1))
                    except:
                        font_size = None

            # Fallback font size based on field height
            if not font_size or font_size <= 0:
                font_size = max(6, min(12, height * 0.65))

            # Use Courier for comb alignment
            c.setFont(default_font, font_size)

            maxlen = annot.get("/MaxLen")
            try:
                maxlen = int(maxlen) if maxlen else None
            except:
                maxlen = None

            y_center = (y0 + y1) / 2.0
            # Small baseline adjustment
            y_text = y_center - (font_size * 0.35)

            if maxlen and maxlen > 0:
                # Comb-style: place each character into a box segment
                box_w = width / maxlen
                # Only draw up to maxlen chars
                val = val[:maxlen]
                for i, ch in enumerate(val):
                    x_center = x0 + (i * box_w) + (box_w / 2.0)
                    c.drawCentredString(x_center, y_text, ch)
            else:
                # Non-comb: draw left aligned, clipped by width (basic)
                c.drawString(x0 + 2, y_text, val)

        c.save()
        packet.seek(0)

        overlay_reader = PdfReader(packet)
        overlay_page = overlay_reader.pages[0]
        page.merge_page(overlay_page)

    # After flattening, remove form fields so viewers don't rely on appearances
    for page in writer.pages:
        if "/Annots" in page:
            del page["/Annots"]

    if "/AcroForm" in writer._root_object:
        del writer._root_object["/AcroForm"]

def fill_pdf_acroform(template_path: str, field_values: dict) -> bytes:
    reader = PdfReader(template_path)
    writer = PdfWriter()

    for page in reader.pages:
        writer.add_page(page)

    root = reader.trailer["/Root"]
    if "/AcroForm" in root:
        writer._root_object.update({NameObject("/AcroForm"): root["/AcroForm"]})
        acroform = writer._root_object["/AcroForm"]
        acroform.update({NameObject("/NeedAppearances"): BooleanObject(True)})

    # Fill values (keeps your existing logic)
    for page in writer.pages:
        writer.update_page_form_field_values(page, field_values)

    # Your comb autosize (keep as-is)
    autosize_all_combed_fields(writer, padding=0.62, font_name="Courier")

    # ✅ The real fix: flatten filled text onto the page
    flatten_filled_fields(writer, field_values, default_font="Courier")

    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


# =========================
# Field mappers
# =========================
def map_common_fields(user_type: str, data: dict) -> dict:
    fields = {}

    if user_type == "Personal":
        fields["IDNUMBER"] = data.get("id_number", "")
        fields["Surname"] = data.get("surname", "")
        fields["Name"] = data.get("first_name", "")
        fields["Initials"] = make_initials(data.get("first_name", ""))
        fields["Email"] = data.get("email", "")
        fields["Address_line_1"] = data.get("address_line_1", "")
        fields["Address_line_2"] = data.get("address_line_2", "")
        fields["City"] = data.get("city", "")
        fields["BeID"] = data.get("bearer_id", "")
        fields["BeSurname"] = data.get("bearer_surname", "")
        fields["BeInti"] = make_initials(data.get("bearer_name", ""))
        fields["Postal_code"] = data.get("postal_code", "")
    else:
        
        fields["IDNUMBER"] = data.get("company_ck", "")
        fields["Surname"] = data.get("company_name", "")
        fields["Email"] = data.get("email", "")
        fields["Address_line_1"] = data.get("address_line_1", "")
        fields["Address_line_2"] = data.get("address_line_2", "")
        fields["City"] = data.get("city", "")
        fields["BeID"] = data.get("bearer_id", "")
        fields["BeSurname"] = data.get("bearer_surname", "")
        fields["BeInti"] = make_initials(data.get("bearer_name", ""))
        fields["Postal_code"] = data.get("postal_code", "")

        # Representative
        fields["RepID"] = data.get("rep_id_number", "")
        fields["RepSurname"] = data.get("rep_surname", "")
        fields["RepInti"] = make_initials(data.get("rep_name", ""))

    # Vehicle
    fields["LicNumber"] = data.get("license_number", "")
    fields["RegNum"] = data.get("register_number", "")
    fields["VIN"] = data.get("vin", "")
    fields["VehBrand"] = data.get("make", "")
    fields["VehSeries"] = data.get("series_name", "")

    return fields


def map_alv_fields(user_type: str, data: dict) -> dict:
    return map_common_fields(user_type, data)


def map_rlv_fields(user_type: str, data: dict) -> dict:
    return map_common_fields(user_type, data)


def map_nco_fields(user_type: str, data: dict) -> dict:
    return map_common_fields(user_type, data)


# =========================
# Forms registry
# =========================
FORM_REGISTRY = {
    "Vehicle licensing": {
        "template": "forms/ALV_Form.pdf",
        "file_suffix": "ALV_form",
        "mapper": map_alv_fields,
    },
    "Registration of new vehicle": {
        "template": "forms/RLV_form.pdf",
        "file_suffix": "RLV_form",
        "mapper": map_rlv_fields,
    },
    "Selling of vehicle": {
        "template": "forms/NCO_form.pdf",
        "file_suffix": "NCO_form",
        "mapper": map_nco_fields,
    },
}


# =========================
# Navigation helpers
# =========================
def next_step():
    st.session_state.step += 1


def prev_step():
    st.session_state.step -= 1


def clear_all():
    st.session_state.step = 1
    st.session_state.user_type = "Personal"
    st.session_state.data = {}
    st.session_state.generated_pdf_bytes = None
    st.session_state.generated_pdf_name = None


# =========================
# Form components
# =========================
def personal_fields():
    d = st.session_state.data
    d["id_number"] = st.text_input("ID number", value=d.get("id_number", ""))
    d["first_name"] = st.text_input("First name", value=d.get("first_name", ""))
    d["surname"] = st.text_input("Surname", value=d.get("surname", ""))
    d["email"] = st.text_input("Email", value=d.get("email", ""))
    d["cellphone"] = st.text_input("Cellphone number", value=d.get("cellphone", ""))
    d["address_line_1"] = st.text_input("Address line 1", value=d.get("address_line_1", ""))
    d["address_line_2"] = st.text_input("Address line 2", value=d.get("address_line_2", ""))
    d["city"] = st.text_input("City", value=d.get("city", ""))
    d["postal_code"] = st.text_input("Postal Code", value=d.get("postal_code", ""))


def company_fields():
    d = st.session_state.data
    d["company_ck"] = st.text_input("Company CK number", value=d.get("company_ck", ""))
    d["company_name"] = st.text_input("Company name", value=d.get("company_name", ""))
    d["rep_id_number"] = st.text_input("Representative ID number", value=d.get("rep_id_number", ""))
    d["rep_name"] = st.text_input("Representative Name", value=d.get("rep_name", ""))
    d["rep_surname"] = st.text_input("Representative Surname", value=d.get("rep_surname", ""))
    d["address_line_1"] = st.text_input("Address line 1", value=d.get("address_line_1", ""))
    d["address_line_2"] = st.text_input("Address line 2", value=d.get("address_line_2", ""))
    d["city"] = st.text_input("City", value=d.get("city", ""))
    d["postal_code"] = st.text_input("Postal Code", value=d.get("postal_code", ""))
    d["email"] = st.text_input("Email", value=d.get("email", ""))
    d["cellphone"] = st.text_input("Cellnumber", value=d.get("cellphone", ""))


def vehicle_fields():
    d = st.session_state.data
    d["license_number"] = st.text_input("License Number", value=d.get("license_number", ""))
    d["register_number"] = st.text_input("Register Number", value=d.get("register_number", ""))
    d["vin"] = st.text_input("VIN", value=d.get("vin", ""))
    d["make"] = st.text_input("Make", value=d.get("make", ""))
    d["series_name"] = st.text_input("Series name", value=d.get("series_name", ""))

def bearer_fields():
    d = st.session_state.data
    st.markdown("#### Person doing the transaction")
    d["bearer_id"] = st.text_input("Bearer ID number", value=d.get("bearer_id", ""))
    d["bearer_name"] = st.text_input("Bearer name", value=d.get("bearer_name", ""))
    d["bearer_surname"] = st.text_input("Bearer surname", value=d.get("bearer_surname", ""))



# =========================
# Validation
# =========================
def is_blank(value) -> bool:
    return value is None or str(value).strip() == ""


def validate_step(step: int) -> tuple[bool, str]:
    d = st.session_state.data

    if step == 1:
        return True, ""

    if step == 2:
        if st.session_state.user_type == "Personal":
            required = ["id_number", "first_name", "surname", "email", "cellphone", "address_line_1", "city", "postal_code"]
        else:
            required = ["company_ck", "company_name", "rep_id_number", "rep_name", "rep_surname", "address_line_1", "city", "postal_code", "email", "cellphone"]

        missing = [k for k in required if is_blank(d.get(k, ""))]
        if missing:
            return False, "Please complete: " + ", ".join(missing)
        return True, ""

    if step == 3:
        if is_blank(d.get("transaction_type", "")):
            return False, "Please select a transaction type."
        return True, ""

    if step == 4:
        required = ["license_number", "register_number", "vin", "make", "series_name"]
        missing = [k for k in required if is_blank(d.get(k, ""))]
        if missing:
            return False, "Please complete: " + ", ".join(missing)
        return True, ""
    
    if step == 5:
        required = ["bearer_id", "bearer_name", "bearer_surname"]
        missing = [k for k in required if is_blank(d.get(k, ""))]
        if missing:
            return False, "Please complete: " + ", ".join(missing)
        return True, ""

    return True, ""


# =========================
# Review UI
# =========================
def pretty_review(user_type: str, data: dict):
    st.subheader("Step 6: Review")

    # -------------------------
    # View details
    # -------------------------
    with st.expander("View details", expanded=True):
        st.json(data)

    st.divider()

    # -------------------------
    # Download captured JSON
    # -------------------------
    st.download_button(
        "Download captured data (JSON)",
        data=json.dumps(make_json_safe({"applicant_type": user_type, **data}), indent=2),
        file_name="captured_data.json",
        mime="application/json",
    )

    st.divider()

    # -------------------------
    # Generate combined PDF (Form + Permission Letter)
    # -------------------------
    tx = data.get("transaction_type", "")
    if tx not in FORM_REGISTRY:
        st.info("No PDF template is registered for this transaction yet.")
        return

    cfg = FORM_REGISTRY[tx]

    st.subheader("Generate your documents")
    col_a, col_b = st.columns([1, 1])

    with col_a:
        if st.button(f"Generate PDF Pack: {tx}"):
            try:
                # 1) Vehicle form PDF
                field_values = cfg["mapper"](user_type, data)
                form_pdf_bytes = fill_pdf_acroform(cfg["template"], field_values)

                # 2) Permission letter PDF
                permission_pdf_bytes = make_permission_letter_pdf(user_type, data)

                # 3) Merge into one PDF
                combined_pdf_bytes = merge_pdfs(form_pdf_bytes, permission_pdf_bytes)

                # 4) Build filename
                if user_type == "Personal":
                    initial = (data.get("first_name", "")[:1] or "X").upper()
                    surname = data.get("surname", "Unknown")
                else:
                    # Company: use representative details
                    initial = (data.get("rep_name", "")[:1] or "X").upper()
                    surname = data.get("rep_surname", "Unknown")

                # Make filename safe
                initial = "".join(c for c in initial if c.isalnum())
                surname = "".join(c for c in surname if c.isalnum())

                filename = f"{initial}_{surname}_{cfg['file_suffix']}.pdf"

                # Save into session state
                st.session_state.generated_pdf_bytes = combined_pdf_bytes
                st.session_state.generated_pdf_name = filename

                st.success("PDF pack generated successfully (Form + Permission Letter).")
                st.rerun()

            except FileNotFoundError:
                st.error(f"Template not found: {cfg['template']}")
            except Exception as e:
                st.error(f"Could not generate the PDF pack: {e}")

    with col_b:
        if st.session_state.generated_pdf_bytes and st.session_state.generated_pdf_name:
            st.download_button(
                "Download PDF Pack",
                data=st.session_state.generated_pdf_bytes,
                file_name=st.session_state.generated_pdf_name,
                mime="application/pdf",
            )
        else:
            st.caption("Click Generate first. The download will appear here.")

    st.divider()

    # -------------------------
    # Preview
    # -------------------------
    st.subheader("Preview")
    if st.session_state.generated_pdf_bytes:
        pdf_preview(st.session_state.generated_pdf_bytes, height=950)
    else:
        st.info("Generate the PDF pack to preview it here.")

    # -------------------------
    # Developer helper (optional)
    # -------------------------
    with st.expander("Developer: Show PDF field names"):
        try:
            st.write(list_pdf_fields(cfg["template"]))
        except Exception as e:
            st.write(f"Could not read fields: {e}")




# =========================
# Main UI
# =========================
st.title("Vehicle Form Assistant (Wizard)")

with st.sidebar:
    st.header("Controls")
    if st.button("Start over"):
        clear_all()
        st.rerun()

steps_total = 6
st.progress((st.session_state.step - 1) / (steps_total - 1))
st.caption(f"Step {st.session_state.step} of {steps_total}")

if st.session_state.step == 1:
    st.subheader("Step 1: Applicant Type")
    st.session_state.user_type = st.radio(
        "Is this for Personal or Company?",
        ["Personal", "Company"],
        index=0 if st.session_state.user_type == "Personal" else 1,
    )

elif st.session_state.step == 2:
    st.subheader("Step 2: Applicant Details")
    if st.session_state.user_type == "Personal":
        personal_fields()
    else:
        company_fields()

elif st.session_state.step == 3:
    st.subheader("Step 3: Transaction Type")
    d = st.session_state.data
    current = d.get("transaction_type", "")
    options = [""] + list(FORM_REGISTRY.keys())
    idx = options.index(current) if current in options else 0
    d["transaction_type"] = st.selectbox("Select the transaction type", options, index=idx)

elif st.session_state.step == 4:
    st.subheader("Step 4: Vehicle Details")
    vehicle_fields()

elif st.session_state.step == 5:
    st.subheader("Step 5: Person doing the transaction")
    bearer_fields()

elif st.session_state.step == 6:
    pretty_review(st.session_state.user_type, st.session_state.data)

st.divider()
left, right = st.columns(2)



with left:
    if st.session_state.step > 1:
        if st.button("Back"):
            prev_step()
            st.rerun()

with right:
    if st.session_state.step < steps_total:
        if st.button("Next"):
            ok, msg = validate_step(st.session_state.step)
            if not ok:
                st.error(msg)
            else:
                st.session_state.generated_pdf_bytes = None
                st.session_state.generated_pdf_name = None
                next_step()
                st.rerun()
    else:
        if st.button("Finish (reset)"):
            clear_all()
            st.rerun()
